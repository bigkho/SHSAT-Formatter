from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

copyPasteDoc = raw_input("Enter file name: ")
NEW = Document()

test = open(copyPasteDoc,"r")
lines = test.readlines()
#FIND TITLE
TITLE = ""
title = False
for line in lines:
    for c in line:
        if c=='}' and title:
            title = False
            break
        if title:
            TITLE+=c
        if c=='{':
            title = True
#WRITE OUT TITLE
docTITLE = NEW.add_paragraph()
docTITLE.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
docTITLE.paragraph_format.space_after = Pt(6)
docTITLE.style.font.name = 'Garamond'
docTITLE.style.font.size = Pt(16)
docTITLE.add_run(TITLE).bold=True

#FIND AUTHOR
AUTHOR = ""
author = False
for line in lines:
    for c in line:
        if c==']' and author:
            author = False
            break
        if author:
            AUTHOR+=c
        if c=='[':
            author = True
#WRITE OUT AUTHOR
docAUTHOR = NEW.add_paragraph()
docAUTHOR.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
docAUTHOR.paragraph_format.space_after = Pt(6)
docAUTHOR.style.font.size = Pt(12)
docAUTHOR.add_run("By "+AUTHOR).italic=True

#PRINT OUT ALL PARAGRAPHS
write = False
paragraph = ""
counter = 0
paragraph_start = False
for line in lines:
    line = line.rstrip()
    line = ''.join(c for c in line if valid_xml_char_ordinal(c))
    print(line)
    if write and line!="PARAGRAPH" and line!="STOP":
        if paragraph_start:
            docPARA.add_run(line)
        else:
            docPARA.add_run(" "+line)
        paragraph_start = False
    if line=="PARAGRAPH":
        docPARA = NEW.add_paragraph()
        docPARA.paragraph_format.line_spacing = 1
        docPARA.paragraph_format.space_after = Pt(6)
        #if counter !=0:
        #    docPARA.add_run("\n")
        counter+=1
        write = True
        paragraph_start = True
    if line=="STOP":
        write = False
        break

#PRINT OUT QUESTIONS
counter = 0
questionNum = 0
question = False
choiceA = False
choiceB = False
choiceC = False
choiceD = False
write = False
for line in lines:
    line = line.rstrip()
    line = ''.join(c for c in line if valid_xml_char_ordinal(c))
    print(line)
    if write and line!="STOP":
        if (line.isdigit()):
            choiceD= False
        if choiceD:
            docPARA.add_run(" "+line)
        if (line[:2]=="D)"):
            choiceC = False
            docPARA = NEW.add_paragraph()
            docPARA.paragraph_format.space_after = Pt(11)
            docPARA.add_run("             D. ").bold=True
            docPARA.add_run(line[2:])
            choiceD = True
        if choiceC:
            docPARA.add_run(" "+line)
        if (line[:2]=="C)"):
            choiceB = False
            docPARA = NEW.add_paragraph()
            docPARA.paragraph_format.space_after = Pt(0)
            docPARA.add_run("             C. ").bold=True
            docPARA.add_run(line[2:])
            choiceC = True
        if choiceB:
            docPARA.add_run(" "+line)
        if (line[:2]=="B)"):
            choiceA = False
            docPARA = NEW.add_paragraph()
            docPARA.paragraph_format.space_after = Pt(0)
            docPARA.add_run("             B. ").bold=True
            docPARA.add_run(line[2:])
            choiceB = True
        if choiceA:
            docPARA.add_run(" "+line)
        if (line[:2]=="A)"):
            question = False
            docPARA = NEW.add_paragraph()
            docPARA.paragraph_format.space_after = Pt(0)
            docPARA.add_run("             A. ").bold=True
            docPARA.add_run(line[2:])
            choiceA = True
        if question:
            docPARA.add_run(" "+line)
        if (line.isdigit()):
            docPARA = NEW.add_paragraph()
            docPARA.bold = True
            docPARA.paragraph_format.space_after = Pt(6)
            question = True
            questionNum += 1
            docPARA.add_run("      "+str(questionNum)+".   ").bold = True

    if line=="STOP":
        write = True
    counter+=1

sections = NEW.sections
for section in sections:
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

NEW.save('NEW.docx')
#passageTest.txt
